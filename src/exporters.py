"""Export to PDF, DOCX, and XLSX formats."""

import logging
import tempfile
from pathlib import Path
from typing import List, Optional
from datetime import datetime

from reportlab.lib.pagesizes import A4, letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, Image as RLImage, PageBreak, Table, TableStyle
)
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib import colors

try:
    import qrcode
    QRCODE_AVAILABLE = True
except ImportError:
    QRCODE_AVAILABLE = False
    logging.warning("qrcode not available, QR codes will be skipped")

try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logging.warning("python-docx not available")

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False
    logging.warning("openpyxl not available")

logger = logging.getLogger(__name__)


class PDFExporter:
    """Export to PDF."""
    
    def __init__(self):
        """Initialize PDF exporter."""
        self.styles = getSampleStyleSheet()
        self._setup_styles()
    
    def _setup_styles(self):
        """Setup custom styles."""
        self.styles.add(ParagraphStyle(
            name='Title',
            parent=self.styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#1f4788'),
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        ))
        self.styles.add(ParagraphStyle(
            name='ProductTitle',
            parent=self.styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#2e5c8a'),
            spaceAfter=12,
            fontName='Helvetica-Bold'
        ))
    
    def export(self, houses: List, output_path: Path) -> None:
        """Export to PDF."""
        output_path = Path(output_path)
        logger.info(f"Generating PDF: {output_path}")
        
        try:
            doc = SimpleDocTemplate(
                str(output_path),
                pagesize=A4,
                rightMargin=0.5*inch,
                leftMargin=0.5*inch,
                topMargin=0.5*inch,
                bottomMargin=0.5*inch
            )
            
            story = []
            
            # Title page
            story.append(Paragraph("Mini House Catalog", self.styles['Title']))
            story.append(Spacer(1, 0.3*inch))
            story.append(Paragraph(
                f"Generated: {datetime.now().strftime('%d.%m.%Y %H:%M')}",
                self.styles['Normal']
            ))
            story.append(Spacer(1, 0.2*inch))
            story.append(Paragraph(
                f"Total Products: {len(houses)}",
                self.styles['Normal']
            ))
            story.append(PageBreak())
            
            # Products
            for idx, house in enumerate(houses, 1):
                # Title with price
                title = f"{house.title}"
                if house.price:
                    title += f" - {house.price:,.0f} RUB"
                story.append(Paragraph(title, self.styles['ProductTitle']))
                story.append(Spacer(1, 0.1*inch))
                
                # Info table
                info_data = []
                if house.size:
                    info_data.append(['Size:', house.size])
                if house.insulation:
                    info_data.append(['Insulation:', house.insulation])
                if house.phone:
                    info_data.append(['Phone:', house.phone])
                if house.email:
                    info_data.append(['Email:', house.email])
                
                if info_data:
                    table = Table(info_data, colWidths=[1.5*inch, 4.5*inch])
                    table.setStyle(TableStyle([
                        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#f0f0f0')),
                        ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
                        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
                        ('FONTSIZE', (0, 0), (-1, -1), 10),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                        ('GRID', (0, 0), (-1, -1), 1, colors.grey)
                    ]))
                    story.append(table)
                    story.append(Spacer(1, 0.15*inch))
                
                # Images
                if house.images:
                    max_width = 6*inch
                    max_height = 2.5*inch
                    for img_path in house.images[:3]:
                        if isinstance(img_path, Path) and img_path.exists():
                            try:
                                rl_image = RLImage(str(img_path), width=max_width, height=max_height)
                                story.append(rl_image)
                                story.append(Spacer(1, 0.1*inch))
                            except Exception as e:
                                logger.debug(f"Could not add image: {e}")
                
                # QR Code
                if QRCODE_AVAILABLE and house.url:
                    try:
                        qr = qrcode.QRCode(
                            version=1,
                            error_correction=qrcode.constants.ERROR_CORRECT_L,
                            box_size=5,
                            border=2
                        )
                        qr.add_data(house.url)
                        qr.make(fit=True)
                        qr_img = qr.make_image(fill_color="black", back_color="white")
                        
                        # Save to temp
                        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                            qr_img.save(tmp.name)
                            rl_qr = RLImage(tmp.name, width=1*inch, height=1*inch)
                            story.append(rl_qr)
                    except Exception as e:
                        logger.debug(f"Could not generate QR: {e}")
                
                story.append(Spacer(1, 0.2*inch))
                story.append(PageBreak())
            
            doc.build(story)
            logger.info(f"✓ PDF exported successfully")
        
        except Exception as e:
            logger.error(f"Error exporting PDF: {e}")
            raise


class DocxExporter:
    """Export to DOCX."""
    
    def export(self, houses: List, output_path: Path) -> None:
        """Export to DOCX."""
        if not DOCX_AVAILABLE:
            logger.error("python-docx is not installed")
            return
        
        output_path = Path(output_path)
        logger.info(f"Generating DOCX: {output_path}")
        
        try:
            doc = Document()
            doc.add_heading('Mini House Catalog', 0)
            doc.add_paragraph(f"Generated: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
            doc.add_paragraph(f"Total Products: {len(houses)}")
            doc.add_paragraph()
            
            for idx, house in enumerate(houses, 1):
                # Title
                title = f"{house.title}"
                if house.price:
                    title += f" - {house.price:,.0f} RUB"
                doc.add_heading(title, level=2)
                
                # Info
                if house.size:
                    p = doc.add_paragraph()
                    p.add_run('Size: ').bold = True
                    p.add_run(house.size or '')
                if house.insulation:
                    p = doc.add_paragraph()
                    p.add_run('Insulation: ').bold = True
                    p.add_run(house.insulation or '')
                if house.phone:
                    p = doc.add_paragraph()
                    p.add_run('Phone: ').bold = True
                    p.add_run(house.phone or '')
                if house.email:
                    p = doc.add_paragraph()
                    p.add_run('Email: ').bold = True
                    p.add_run(house.email or '')
                
                doc.add_paragraph()
                
                # Images
                if house.images:
                    for img_path in house.images[:3]:
                        if isinstance(img_path, Path) and img_path.exists():
                            try:
                                doc.add_picture(str(img_path), width=Inches(5))
                            except Exception as e:
                                logger.debug(f"Could not add image: {e}")
                
                # QR Code
                if QRCODE_AVAILABLE and house.url:
                    try:
                        qr = qrcode.QRCode(
                            version=1,
                            error_correction=qrcode.constants.ERROR_CORRECT_L,
                            box_size=5,
                            border=2
                        )
                        qr.add_data(house.url)
                        qr.make(fit=True)
                        qr_img = qr.make_image(fill_color="black", back_color="white")
                        
                        # Save to temp
                        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                            qr_img.save(tmp.name)
                            doc.add_picture(tmp.name, width=Inches(1.5))
                    except Exception as e:
                        logger.debug(f"Could not generate QR: {e}")
                
                doc.add_page_break()
            
            doc.save(str(output_path))
            logger.info(f"✓ DOCX exported successfully")
        
        except Exception as e:
            logger.error(f"Error exporting DOCX: {e}")
            raise


class ExcelExporter:
    """Export to XLSX."""
    
    def export(self, houses: List, output_path: Path) -> None:
        """Export to XLSX."""
        if not XLSX_AVAILABLE:
            logger.error("openpyxl is not installed")
            return
        
        output_path = Path(output_path)
        logger.info(f"Generating XLSX: {output_path}")
        
        try:
            wb = Workbook()
            ws = wb.active
            ws.title = "Catalog"
            
            # Headers
            headers = ['№', 'Title', 'Price (RUB)', 'Size', 'Insulation', 'Phone', 'Email', 'URL']
            ws.append(headers)
            
            # Style header
            header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF")
            for cell in ws[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center", vertical="center")
            
            # Data
            for idx, house in enumerate(houses, 1):
                row = [
                    idx,
                    house.title or "",
                    house.price or "",
                    house.size or "",
                    house.insulation or "",
                    house.phone or "",
                    house.email or "",
                    house.url or ""
                ]
                ws.append(row)
            
            # Adjust column widths
            ws.column_dimensions['A'].width = 5
            ws.column_dimensions['B'].width = 30
            ws.column_dimensions['C'].width = 15
            ws.column_dimensions['D'].width = 15
            ws.column_dimensions['E'].width = 15
            ws.column_dimensions['F'].width = 15
            ws.column_dimensions['G'].width = 20
            ws.column_dimensions['H'].width = 40
            
            wb.save(str(output_path))
            logger.info(f"✓ XLSX exported successfully")
        
        except Exception as e:
            logger.error(f"Error exporting XLSX: {e}")
            raise
